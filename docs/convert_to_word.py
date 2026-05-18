import os
import re
import sys

def main():
    try:
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("python-docx is not installed. Installing it now...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

    # Paths
    base_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base_dir, "business_whitepaper.md")
    docx_path = os.path.join(base_dir, "business_whitepaper.docx")

    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return

    print("Generating beautifully styled Word Document...")
    
    # Initialize Document
    doc = docx.Document()

    # Set Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style Configurations
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Custom Helper for Shaded Callout Box/Table Borders
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Skip center badges and logo wrappers
        if stripped.startswith("<div") or stripped.startswith("</div") or stripped.startswith("!["):
            continue

        # Handle Tables
        if stripped.startswith("|"):
            if "---" in stripped:
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            continue
        else:
            if in_table:
                # Compile and write table
                num_cols = len(table_headers)
                if num_cols > 0:
                    t = doc.add_table(rows=1, cols=num_cols)
                    t.style = 'Light Shading Accent 1'
                    hdr_cells = t.rows[0].cells
                    for i in range(num_cols):
                        hdr_cells[i].text = table_headers[i]
                        set_cell_background(hdr_cells[i], "7F55F7")  # Deep Purple Primary Brand
                        # White text for header
                        for p in hdr_cells[i].paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                r.font.name = 'Arial'
                    
                    for row_data in table_rows:
                        row_cells = t.add_row().cells
                        for i in range(min(num_cols, len(row_data))):
                            row_cells[i].text = row_data[i]
                            # Clean custom styling for runs
                            for p in row_cells[i].paragraphs:
                                for r in p.runs:
                                    r.font.name = 'Arial'
                                    r.font.size = Pt(10)
                                    # Highlight bold tags
                                    if "**" in r.text:
                                        r.text = r.text.replace("**", "")
                                        r.font.bold = True

                # Reset table state
                in_table = False
                table_headers = []
                table_rows = []
                doc.add_paragraph()  # Space
                continue

        # Handle Headers
        if stripped.startswith("#"):
            match = re.match(r'^(#+)\s*(.*)$', stripped)
            if match:
                level = len(match.group(1))
                title = match.group(2)
                
                # Strip markdown bold formatting inside titles
                title = title.replace("**", "").replace("`", "")

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True

                run = p.add_run(title)
                run.font.name = 'Arial'
                run.font.bold = True

                if level == 1:
                    run.font.size = Pt(20)
                    run.font.color.rgb = RGBColor(0x7F, 0x55, 0xF7)  # Primary Purple
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif level == 2:
                    run.font.size = Pt(15)
                    run.font.color.rgb = RGBColor(0x00, 0x96, 0x88)  # Secondary Teal
                else:
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Handle Quote Blocks / Callouts
        if stripped.startswith(">"):
            content = stripped.lstrip(">").strip()
            # Skip highlight tags
            if content.startswith("[!"):
                continue
            
            # Draw as a single shaded callout cell table
            t = doc.add_table(rows=1, cols=1)
            cell = t.cell(0, 0)
            set_cell_background(cell, "F3EFFF")  # Soft Lavender background
            
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # Clean content markdown bold/italic tags
            content = content.replace("**", "").replace("*", "").replace("`", "")
            run = p.add_run(content)
            run.font.name = 'Arial'
            run.font.italic = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x55, 0x33, 0xAA)
            doc.add_paragraph()  # Space
            continue

        # Handle Bullet Points
        if stripped.startswith("*") or stripped.startswith("-"):
            content = stripped[1:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            
            # Handle inline bold elements
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    clean_part = part[2:-2]
                    r = p.add_run(clean_part)
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Arial'
            continue

        # Handle Text Blocks
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
            # Handle inline bold elements
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    clean_part = part[2:-2]
                    r = p.add_run(clean_part)
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Arial'

    # Save Document
    doc.save(docx_path)
    print(f"Successfully generated {docx_path}!")

if __name__ == "__main__":
    main()
